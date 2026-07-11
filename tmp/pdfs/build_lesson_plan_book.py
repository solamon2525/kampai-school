import re
import sys
from pathlib import Path

try:
    sys.stdout.reconfigure(encoding="utf-8")
except Exception:
    pass

sys.path.insert(0, str(Path(r"D:\School คำไผ่\kampai-school\tmp\pdfs\pymupdf")))

import pymupdf
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


SOURCE = Path(r"D:\แผนจาก DLTV\คู่มือครูและแผนการสอน_คณิตศาสตร์_ป.4_ภาคเรียนที่_1-2569-07112009.pdf")
OUT_DIR = Path(r"D:\School คำไผ่\kampai-school\output\documents")
OUT = OUT_DIR / "แผนการสอน_คณิตศาสตร์_ป4_ภาคเรียนที่1_2569.docx"

NAVY = "17365D"
GOLD = "C59A32"
PALE_GOLD = "FFF7E1"
PALE_BLUE = "EEF3F8"
INK = "1F2937"
MUTED = "667085"
WHITE = "FFFFFF"
CREAM = "FFF7E1"
BORDER = "CBD5E1"
EMBLEM = Path(r"D:\School คำไผ่\kampai-school\tmp\assets\emblem.png")
# Brand font. Sarabun is installed per-user and embedded into the exported docx/PDF
# so glyphs render consistently regardless of the reader's machine.
FONT = "Sarabun"


def clean(text):
    text = text.replace("\uf06f", "□").replace("\u200b", " ")
    text = re.sub(r"\s+", " ", text).strip()
    return text


def shorten(text, limit):
    text = clean(text)
    if len(text) <= limit:
        return text
    cut = text[:limit].rsplit(" ", 1)[0]
    return cut.rstrip(" ,;:") + "…"


def extract_between(text, start_patterns, end_patterns, limit=420):
    start = None
    for pat in start_patterns:
        m = re.search(pat, text, re.S)
        if m:
            start = m.end()
            break
    if start is None:
        return ""
    tail = text[start:]
    end = len(tail)
    for pat in end_patterns:
        m = re.search(pat, tail, re.S)
        if m:
            end = min(end, m.start())
    return shorten(tail[:end], limit)


def detect_plans(doc):
    plans = []
    current_unit = 1
    seen = set()
    for idx in range(54, len(doc)):
        text = doc[idx].get_text()
        if idx + 1 >= 1515:
            current_unit = 2
        m = re.search(r"แผนการจัดการเรียนรู้\s*ที่\s*(\d+)\s*(?:เรื่อง)?\s*([^\n]*)", text)
        if not m:
            continue
        number = int(m.group(1))
        if current_unit == 1 and not 1 <= number <= 62:
            continue
        if current_unit == 2 and not 1 <= number <= 10:
            continue
        key = (current_unit, number)
        if key in seen:
            continue
        title = clean(m.group(2))
        if not title or len(title) < 4:
            continue
        seen.add(key)
        plans.append({"unit": current_unit, "number": number, "title": title, "source_page": idx + 1})
    plans.sort(key=lambda x: (x["unit"], x["number"]))
    if len(plans) != 72:
        raise RuntimeError(f"Expected 72 plans, found {len(plans)}")
    for pos, plan in enumerate(plans):
        start = plan["source_page"] - 1
        # The overview normally occupies two pages. Include the first teaching-grid page
        # so that objectives and the opening activity are available for summarisation.
        text = "\n".join(doc[i].get_text() for i in range(start, min(start + 4, len(doc))))
        plan["raw"] = text
        plan["standard"] = extract_between(
            text,
            [r"1\.\s*มาตรฐานการเรียนรู้/ตัวชี้วัด", r"มาตรฐานการเรียนรู้/ตัวชี้วัด"],
            [r"ตัวชี้วัด\s*ป", r"2\.\s*สาระสำคัญ", r"สาระสำคัญ/ความคิดรวบยอด"],
            260,
        )
        plan["indicator"] = extract_between(
            text,
            [r"ตัวชี้วัด\s*(?=ป\s*\.?\s*[๔4])", r"ตัวชี้วัด\s*"],
            [r"2\.\s*สาระสำคัญ", r"สาระสำคัญ/ความคิดรวบยอด"],
            220,
        )
        plan["content"] = extract_between(
            text,
            [r"3\.\s*สาระการเรียนรู้"],
            [r"4\.\s*จุดประสงค์", r"จุดประสงค์การเรียนรู้"],
            240,
        )
        plan["concept"] = extract_between(
            text,
            [r"2\.\s*สาระสำคัญ/ความคิดรวบยอด", r"2\.\s*สาระสำคัญ", r"สาระสำคัญ/ความคิดรวบยอด"],
            [r"3\.\s*สาระการเรียนรู้", r"4\.\s*จุดประสงค์"],
            260,
        )
        plan["objectives"] = extract_between(
            text,
            [r"4\.\s*จุดประสงค์การเรียนรู้", r"จุดประสงค์การเรียนรู้"],
            [r"5\.\s*สาระการเรียนรู้", r"5\.\s*สมรรถนะ", r"6\.\s*คุณลักษณะ", r"7\.\s*กิจกรรม"],
            390,
        )
        if not plan["standard"]:
            plan["standard"] = "มาตรฐาน ค 1.1 ใช้ความรู้และทักษะกระบวนการทางคณิตศาสตร์ในการแก้ปัญหา"
        if not plan["indicator"]:
            plan["indicator"] = f"อธิบายและแสดงวิธีหาคำตอบเรื่อง {plan['title']} ได้อย่างถูกต้อง"
        if not plan["content"]:
            plan["content"] = plan["title"]
        if not plan["concept"]:
            plan["concept"] = f"เรียนรู้และประยุกต์ใช้เรื่อง {plan['title']} อย่างเป็นขั้นตอนและสมเหตุสมผล"
        if not plan["objectives"]:
            plan["objectives"] = f"อธิบายหลักการและแสดงวิธีหาคำตอบเรื่อง {plan['title']} พร้อมตรวจสอบความสมเหตุสมผลของคำตอบได้"
    return plans


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=70, start=100, bottom=70, end=100):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table, color=BORDER, size="6"):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = borders.find(qn(f"w:{edge}"))
        if tag is None:
            tag = OxmlElement(f"w:{edge}")
            borders.append(tag)
        tag.set(qn("w:val"), "single")
        tag.set(qn("w:sz"), size)
        tag.set(qn("w:color"), color)


def set_table_geometry(table, widths_cm):
    """Make tblW, tblGrid and tcW agree so Word cannot reinterpret columns."""
    widths = [int(cm / 2.54 * 1440) for cm in widths_cm]
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")
    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for cell, width_cm in zip(row.cells, widths_cm):
            set_fixed_width(cell, width_cm)


def set_fixed_width(cell, cm):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    # OOXML table widths use twentieths of a point (DXA), not EMU.
    tc_w.set(qn("w:w"), str(int(cm / 2.54 * 1440)))
    tc_w.set(qn("w:type"), "dxa")


def set_run(run, size=9, bold=False, color=INK, italic=False):
    run.font.name = FONT
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:ascii"), FONT)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    run.font.color.rgb = RGBColor.from_string(color)


def set_para(p, before=0, after=2, line=1.0, align=None, keep=False):
    pf = p.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    pf.line_spacing = line
    if align is not None:
        p.alignment = align
    if keep:
        pf.keep_with_next = True


def clear_para(p):
    for child in list(p._p):
        if child.tag.endswith("}r") or child.tag.endswith("}hyperlink"):
            p._p.remove(child)


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("หน้า ")
    set_run(run, 8, color=MUTED)
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), "PAGE")
    paragraph._p.append(fld)


def add_labeled_text(cell, label, text, label_color=NAVY, size=8.7):
    p = cell.paragraphs[0]
    clear_para(p)
    set_para(p, after=1, line=1.0)
    r = p.add_run(label)
    set_run(r, size, True, label_color)
    r = p.add_run(text)
    set_run(r, size, False, INK)


def add_bullet(cell, text, size=8.5):
    p = cell.add_paragraph()
    set_para(p, after=0.8, line=1.0)
    p.paragraph_format.left_indent = Cm(0.35)
    p.paragraph_format.first_line_indent = Cm(-0.28)
    r = p.add_run("• ")
    set_run(r, size, True, GOLD)
    r = p.add_run(text)
    set_run(r, size)


def set_para_shading(p, fill):
    p_pr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill)
    p_pr.append(shd)


def set_para_left_accent(p, color=GOLD, size="26", space="8"):
    p_pr = p._p.get_or_add_pPr()
    bdr = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), size)
    left.set(qn("w:space"), space)
    left.set(qn("w:color"), color)
    bdr.append(left)
    p_pr.append(bdr)


def add_section_label(doc, text):
    """Accent chip: pale bar + thick gold left rule + navy bold label."""
    container = doc
    p = container.add_paragraph()
    set_para(p, before=5, after=3, line=1.0, keep=True)
    p.paragraph_format.left_indent = Cm(0.28)
    set_para_shading(p, PALE_BLUE)
    set_para_left_accent(p, GOLD)
    r = p.add_run(text)
    set_run(r, 9.6, True, NAVY)
    return p


def add_heading(doc, text, size=18, rule=True):
    """A page/section heading with a gold underline rule."""
    p = doc.add_paragraph()
    set_para(p, before=0, after=2 if rule else 8, line=1.0)
    set_run(p.add_run(text), size, True, NAVY)
    if rule:
        add_hrule(doc, color=GOLD, size="14", before=1, after=10)
    return p


LINE_COLOR = "94A3B8"


def _apply_write_border(p, before, after, indent):
    set_para(p, before=before, after=after, line=1.0)
    p.paragraph_format.left_indent = Cm(indent)
    p_pr = p._p.get_or_add_pPr()
    borders = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "dotted")
    bottom.set(qn("w:sz"), "8")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), LINE_COLOR)
    borders.append(bottom)
    p_pr.append(borders)
    return p


def add_write_line(parent, before=10, after=11, indent=0.1, reuse=None):
    """A dotted rule that acts as a handwriting line on the printed page."""
    p = reuse if reuse is not None else parent.add_paragraph()
    return _apply_write_border(p, before, after, indent)


def add_hrule(doc, color=GOLD, size="12", before=4, after=8, align=None):
    p = doc.add_paragraph()
    set_para(p, before=before, after=after, line=1.0, align=align)
    p_pr = p._p.get_or_add_pPr()
    borders = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), color)
    borders.append(bottom)
    p_pr.append(borders)
    return p


def finish_header_cell(cell):
    """Give a table header cell breathing room so Thai upper marks aren't clipped."""
    set_cell_margins(cell, top=110, bottom=90)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def activity_text(plan):
    title = plan["title"]
    intro = f"ทบทวนความรู้เดิมด้วยคำถามหรือสถานการณ์ใกล้ตัวที่เชื่อมโยงกับ {title} และแจ้งจุดประสงค์การเรียนรู้"
    if any(k in title for k in ["โจทย์ปัญหา", "สร้างโจทย์"]):
        teach = f"ร่วมกันวิเคราะห์สิ่งที่โจทย์กำหนดและสิ่งที่ต้องหา เลือกวิธีดำเนินการ เขียนประโยคสัญลักษณ์ และแสดงวิธีทำเรื่อง {title}; ฝึกเป็นคู่แล้วตรวจคำตอบร่วมกัน"
    elif any(k in title for k in ["คูณ", "หาร", "บวก", "ลบ", "ประมาณ", "เฉลี่ย"]):
        teach = f"ครูสาธิตแนวคิดด้วยตัวอย่างจากง่ายไปยาก นักเรียนอธิบายเหตุผลและฝึกคำนวณเรื่อง {title} เป็นรายคู่/รายบุคคล พร้อมตรวจสอบความสมเหตุสมผล"
    elif any(k in title for k in ["เวลา", "นาฬิกา", "ตาราง"]):
        teach = f"ใช้สถานการณ์และตารางเวลาในชีวิตประจำวัน ให้นักเรียนอ่าน เปรียบเทียบ บันทึก หรือคำนวณเกี่ยวกับ {title} พร้อมอภิปรายวิธีคิด"
    else:
        teach = f"สำรวจตัวอย่างและสื่อที่กำหนด สรุปหลักการเรื่อง {title} ร่วมกัน แล้วฝึกปฏิบัติเป็นคู่และรายบุคคลโดยครูให้คำแนะนำ"
    close = "นักเรียนอธิบายวิธีคิดและสรุปหลักการด้วยภาษาของตนเอง ครูตรวจความเข้าใจด้วยคำถามสั้นหรือบัตรออกจากห้องเรียน"
    return intro, teach, close


def add_body_text(doc, text, size=8.7, after=3, bullet=False):
    p = doc.add_paragraph()
    set_para(p, after=after, line=1.05)
    if bullet:
        p.paragraph_format.left_indent = Cm(0.35)
        p.paragraph_format.first_line_indent = Cm(-0.28)
        set_run(p.add_run("• "), size, True, GOLD)
    else:
        p.paragraph_format.left_indent = Cm(0.18)
    set_run(p.add_run(text), size)
    return p


def add_assessment_table(doc):
    """Standard measurement matrix: what / method / instrument / criterion."""
    rows = [
        ("ด้านความรู้ (K)", "ตรวจผลงาน / แบบฝึกหัด", "แบบประเมินผลงาน", "ผ่านระดับพอใช้ขึ้นไป"),
        ("ด้านทักษะ/กระบวนการ (P)", "สังเกตพฤติกรรมการเรียนรู้", "แบบประเมินทักษะและกระบวนการทางคณิตศาสตร์", "ผ่านระดับพอใช้ขึ้นไป"),
        ("ด้านคุณลักษณะ (A)", "สังเกตพฤติกรรมการเรียนรู้", "แบบประเมินคุณลักษณะอันพึงประสงค์", "ผ่านระดับพอใช้ขึ้นไป"),
    ]
    headers = ["รายการประเมิน", "วิธีวัด", "เครื่องมือ", "เกณฑ์"]
    widths = [3.2, 3.7, 7.0, 4.0]
    table = doc.add_table(rows=len(rows) + 1, cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    for c, h in enumerate(headers):
        cell = table.cell(0, c)
        set_fixed_width(cell, widths[c])
        set_cell_shading(cell, NAVY)
        p = cell.paragraphs[0]
        clear_para(p)
        set_para(p, after=0, line=1.0)
        set_run(p.add_run(h), 8.3, True, WHITE)
        finish_header_cell(cell)
    for r, row in enumerate(rows, start=1):
        for c, value in enumerate(row):
            cell = table.cell(r, c)
            set_fixed_width(cell, widths[c])
            set_cell_margins(cell, top=70, bottom=70, start=110, end=110)
            p = cell.paragraphs[0]
            clear_para(p)
            set_para(p, after=0, line=1.0)
            set_run(p.add_run(value), 8.2, bold=(c == 0), color=(NAVY if c == 0 else INK))
    set_table_geometry(table, widths)
    set_table_borders(table)


def add_plan_page(doc, plan):
    unit_name = "จำนวนนับ และการบวก การลบ การคูณ การหาร" if plan["unit"] == 1 else "การวัด"

    # --- header card: navy title strip + pale meta strip ---
    card = doc.add_table(rows=2, cols=1)
    card.alignment = WD_TABLE_ALIGNMENT.CENTER
    card.autofit = False
    head, meta = card.cell(0, 0), card.cell(1, 0)
    set_fixed_width(head, 17.9)
    set_fixed_width(meta, 17.9)
    set_cell_shading(head, NAVY)
    set_cell_shading(meta, WHITE)
    set_cell_margins(head, top=120, start=170, bottom=110, end=170)
    set_cell_margins(meta, top=110, start=170, bottom=110, end=170)

    p = head.paragraphs[0]
    clear_para(p)
    set_para(p, after=2, line=1.0)
    set_run(p.add_run(f"แผนการจัดการเรียนรู้ที่ {plan['number']}"), 14, True, WHITE)
    set_run(p.add_run(f"      หน่วยการเรียนรู้ที่ {plan['unit']}"), 10, True, GOLD)
    p = head.add_paragraph()
    set_para(p, after=0, line=1.05)
    set_run(p.add_run(plan["title"]), 11.5, True, CREAM)

    meta_rows = [
        ("รายวิชา", "คณิตศาสตร์ ค14101", "ชั้น / เวลา", "ป.4 / 50 นาที"),
        ("หน่วยการเรียนรู้", unit_name, "ภาคเรียน / ปีการศึกษา", "1 / 2569"),
        ("ผู้สอน", "..................................................", "วันที่สอน", "......... / ......... / ............"),
    ]
    for i, (l1, v1, l2, v2) in enumerate(meta_rows):
        p = meta.paragraphs[0] if i == 0 else meta.add_paragraph()
        clear_para(p)
        set_para(p, after=1.5 if i < len(meta_rows) - 1 else 0, line=1.0)
        p.paragraph_format.tab_stops.add_tab_stop(Cm(9.0))
        set_run(p.add_run(f"{l1}: "), 8.6, True, NAVY)
        set_run(p.add_run(v1), 8.6)
        set_run(p.add_run(f"\t{l2}: "), 8.6, True, NAVY)
        set_run(p.add_run(v2), 8.6)
    set_table_geometry(card, [17.9])
    set_table_borders(card, color=NAVY, size="6")
    doc.add_paragraph().paragraph_format.space_after = Pt(3)

    add_section_label(doc, "มาตรฐานการเรียนรู้")
    add_body_text(doc, shorten(plan["standard"], 250))

    add_section_label(doc, "ตัวชี้วัด")
    add_body_text(doc, shorten(plan["indicator"], 220))

    add_section_label(doc, "สาระสำคัญ / ความคิดรวบยอด")
    add_body_text(doc, shorten(plan["concept"], 245))

    add_section_label(doc, "สาระการเรียนรู้")
    add_body_text(doc, shorten(plan["content"], 240))

    add_section_label(doc, "จุดประสงค์การเรียนรู้ (K / P / A)")
    add_body_text(doc, shorten(plan["objectives"], 330), bullet=True, after=1.5)
    add_body_text(
        doc,
        "สมรรถนะสำคัญ: การคิด การแก้ปัญหา และการสื่อสารทางคณิตศาสตร์  |  คุณลักษณะอันพึงประสงค์: ใฝ่เรียนรู้ ซื่อสัตย์สุจริต และมุ่งมั่นในการทำงาน",
        bullet=True,
    )

    # ---------- page 2 ----------
    doc.add_page_break()

    add_section_label(doc, "กระบวนการจัดการเรียนรู้")
    intro, teach, close = activity_text(plan)
    stages = [("ขั้นนำ 5 นาที", intro), ("ขั้นสอน 35 นาที", teach), ("ขั้นสรุป 10 นาที", close)]
    for stage, detail in stages:
        p = doc.add_paragraph()
        set_para(p, after=1.6, line=1.05)
        p.paragraph_format.left_indent = Cm(0.25)
        set_run(p.add_run(stage + ": "), 8.7, True, NAVY)
        set_run(p.add_run(detail), 8.6)

    add_section_label(doc, "สื่อ / แหล่งเรียนรู้")
    add_body_text(doc, "สื่อการสอน (PowerPoint / บทเรียน DLTV), ใบกิจกรรมและแบบฝึกหัดตามแผน, อุปกรณ์ประกอบการสอน และแหล่งเรียนรู้ในห้องเรียน", after=3)

    add_section_label(doc, "ภาระงาน / ชิ้นงาน")
    add_body_text(doc, f"แบบฝึกหัด / ใบกิจกรรมเรื่อง {shorten(plan['title'], 90)} และการบันทึกผลงานลงในสมุด", after=3)

    add_section_label(doc, "การวัดและประเมินผล")
    add_assessment_table(doc)

    # --- record card: bordered box with navy section strips + writing lines ---
    doc.add_paragraph().paragraph_format.space_after = Pt(3)
    rec = doc.add_table(rows=4, cols=1)
    rec.alignment = WD_TABLE_ALIGNMENT.CENTER
    rec.autofit = False
    for r in range(4):
        set_fixed_width(rec.cell(r, 0), 17.9)

    def strip(cell, text):
        set_cell_shading(cell, NAVY)
        set_cell_margins(cell, top=70, start=180, bottom=60, end=180)
        p = cell.paragraphs[0]
        clear_para(p)
        set_para(p, after=0, line=1.0)
        set_run(p.add_run(text), 9.4, True, WHITE)

    strip(rec.cell(0, 0), "บันทึกผลหลังการจัดการเรียนรู้")
    body = rec.cell(1, 0)
    set_cell_shading(body, WHITE)
    set_cell_margins(body, top=100, start=200, bottom=140, end=200)
    p = body.paragraphs[0]
    clear_para(p)
    set_para(p, after=1, line=1.0)
    set_run(p.add_run("นักเรียนทั้งหมด ......... คน     ผ่านเกณฑ์การประเมิน ......... คน     ควรได้รับการพัฒนาเพิ่มเติม ......... คน"), 8.6)
    p = body.add_paragraph()
    set_para(p, before=2, after=0, line=1.0)
    set_run(p.add_run("ผลการจัดการเรียนรู้ / ปัญหา อุปสรรค / ข้อเสนอแนะเพื่อการพัฒนา"), 8.6, True, NAVY)
    for _ in range(4):
        add_write_line(body)

    strip(rec.cell(2, 0), "ความเห็นของผู้บริหารหรือผู้ที่ได้รับมอบหมาย")
    adm = rec.cell(3, 0)
    set_cell_shading(adm, WHITE)
    set_cell_margins(adm, top=100, start=200, bottom=140, end=200)
    add_write_line(adm, reuse=adm.paragraphs[0])
    add_write_line(adm)
    add_write_line(adm)
    p = adm.add_paragraph()
    set_para(p, before=10, after=0, line=1.0)
    p.paragraph_format.tab_stops.add_tab_stop(Cm(8.5))
    set_run(p.add_run("ลงชื่อ ...................................... ผู้สอน"), 8.6)
    set_run(p.add_run("\tลงชื่อ ...................................... ผู้บริหาร"), 8.6)

    set_table_geometry(rec, [17.9])
    set_table_borders(rec, color=NAVY, size="6")

    p = doc.add_paragraph()
    set_para(p, before=3, after=0, line=1.0, align=WD_ALIGN_PARAGRAPH.RIGHT)
    set_run(p.add_run(f"อ้างอิงต้นฉบับ DLTV หน้า PDF {plan['source_page']} เป็นต้นไป"), 7.2, italic=True, color=MUTED)


def add_cover(doc):
    center = WD_ALIGN_PARAGRAPH.CENTER
    # Emblem
    set_para(doc.add_paragraph(), after=0)
    p = doc.add_paragraph()
    set_para(p, after=5, align=center)
    if EMBLEM.exists():
        p.add_run().add_picture(str(EMBLEM), width=Cm(3.3))
    # School kicker
    p = doc.add_paragraph()
    set_para(p, after=1, align=center)
    set_run(p.add_run("โรงเรียนคำไผ่"), 15, True, GOLD)
    p = doc.add_paragraph()
    set_para(p, after=14, align=center)
    set_run(p.add_run("กลุ่มสาระการเรียนรู้คณิตศาสตร์"), 11.5, False, MUTED)

    # Title block
    p = doc.add_paragraph()
    set_para(p, after=4, align=center)
    set_run(p.add_run("แผนการจัดการเรียนรู้"), 34, True, NAVY)
    p = doc.add_paragraph()
    set_para(p, before=2, after=10, align=center)
    set_run(p.add_run("◆ ◆ ◆"), 12, True, GOLD)
    p = doc.add_paragraph()
    set_para(p, after=6, align=center)
    set_run(p.add_run("รายวิชาคณิตศาสตร์  •  รหัสวิชา ค14101"), 18, True, GOLD)
    p = doc.add_paragraph()
    set_para(p, after=4, align=center)
    set_run(p.add_run("ชั้นประถมศึกษาปีที่ 4"), 15, True, INK)
    p = doc.add_paragraph()
    set_para(p, after=4, align=center)
    set_run(p.add_run("ภาคเรียนที่ 1  ปีการศึกษา 2569"), 14, False, INK)
    p = doc.add_paragraph()
    set_para(p, after=0, align=center)
    set_run(p.add_run("ฉบับเรียบเรียงเพื่อใช้จัดการเรียนรู้และเข้าเล่ม"), 11, italic=True, color=MUTED)

    # Push the identity card toward the lower third
    for _ in range(7):
        set_para(doc.add_paragraph(), after=0)

    # Centered identity card (pale gold, gold border)
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    cell = table.cell(0, 0)
    set_fixed_width(cell, 11.5)
    set_cell_shading(cell, PALE_GOLD)
    set_cell_margins(cell, top=180, start=220, bottom=180, end=220)
    rows = [
        ("ผู้สอน", "..............................................................."),
        ("กลุ่มสาระการเรียนรู้", "คณิตศาสตร์"),
        ("โรงเรียน", "โรงเรียนคำไผ่"),
        ("สังกัด", "..............................................................."),
    ]
    for i, (label, value) in enumerate(rows):
        p = cell.paragraphs[0] if i == 0 else cell.add_paragraph()
        clear_para(p)
        set_para(p, after=4 if i < len(rows) - 1 else 0, line=1.0, align=center)
        set_run(p.add_run(f"{label}   "), 11.5, True, NAVY)
        set_run(p.add_run(value), 11.5, False, INK)
    set_table_geometry(table, [11.5])
    set_table_borders(table, color=GOLD, size="8")

    # Bottom credit line
    for _ in range(6):
        set_para(doc.add_paragraph(), after=0)
    p = doc.add_paragraph()
    set_para(p, before=6, align=center)
    set_run(p.add_run("เรียบเรียงจากชุดกิจกรรมการเรียนรู้ DLTV มูลนิธิการศึกษาทางไกลผ่านดาวเทียม ในพระบรมราชูปถัมภ์"), 9, color=MUTED)


def add_front_matter(doc, plans, lead_break=True):
    if lead_break:
        doc.add_page_break()
    p = doc.add_paragraph()
    set_para(p, after=2, align=WD_ALIGN_PARAGRAPH.CENTER)
    set_run(p.add_run("คำนำ"), 20, True, NAVY)
    p = doc.add_paragraph()
    set_para(p, after=12, align=WD_ALIGN_PARAGRAPH.CENTER)
    set_run(p.add_run("◆ ◆ ◆"), 12, True, GOLD)
    paras = [
        "แผนการจัดการเรียนรู้ฉบับนี้จัดทำขึ้นสำหรับรายวิชาคณิตศาสตร์ ชั้นประถมศึกษาปีที่ 4 ภาคเรียนที่ 1 ปีการศึกษา 2569 โดยคัดสาระสำคัญจากคู่มือครูและแผนการสอนของ DLTV แล้วเรียบเรียงใหม่ให้กระชับ เหมาะสำหรับใช้เตรียมการสอน บันทึกหลักฐาน และจัดพิมพ์เข้าเล่ม",
        "เนื้อหาประกอบด้วย 72 แผน แบ่งเป็นหน่วยที่ 1 จำนวนนับ และการบวก การลบ การคูณ การหาร จำนวน 62 แผน และหน่วยที่ 2 การวัด จำนวน 10 แผน แต่ละแผนกำหนดเวลา 50 นาที พร้อมมาตรฐาน จุดประสงค์ กิจกรรมการเรียนรู้ สื่อ และแนวทางประเมินผล",
        "ผู้สอนควรปรับตัวอย่าง สื่อ และระดับความยากให้สอดคล้องกับบริบทของผู้เรียน ตลอดจนบันทึกผลหลังสอนเพื่อนำไปพัฒนาการจัดการเรียนรู้ในครั้งต่อไป",
    ]
    for text in paras:
        p = doc.add_paragraph()
        set_para(p, after=9, line=1.3, align=WD_ALIGN_PARAGRAPH.LEFT)
        set_run(p.add_run(text), 11)
    p = doc.add_paragraph()
    set_para(p, before=26, align=WD_ALIGN_PARAGRAPH.RIGHT)
    set_run(p.add_run("ผู้จัดทำ\nปีการศึกษา 2569"), 10, color=MUTED)

    doc.add_page_break()
    add_heading(doc, "คำอธิบายรายวิชาและโครงสร้างเวลา", size=18)
    p = doc.add_paragraph()
    set_para(p, after=10, line=1.3, align=WD_ALIGN_PARAGRAPH.LEFT)
    set_run(p.add_run("ศึกษา ฝึกทักษะการอ่าน การเขียน การเปรียบเทียบ และการประมาณค่าจำนวนนับ การคำนวณและแก้โจทย์ปัญหาการบวก การลบ การคูณ การหาร การคำนวณระคน ค่าเฉลี่ย ตลอดจนการบอกและคำนวณเกี่ยวกับเวลา โดยใช้กระบวนการแก้ปัญหา การให้เหตุผล การสื่อสาร และการเชื่อมโยงความรู้กับสถานการณ์ในชีวิตประจำวัน"), 10.5)
    table = doc.add_table(rows=3, cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    widths = [1.4, 9.6, 2.5, 3.7]
    headers = ["หน่วย", "ชื่อหน่วยการเรียนรู้", "จำนวนแผน", "เวลาโดยประมาณ"]
    data = [("1", "จำนวนนับ และการบวก การลบ การคูณ การหาร", "62", "51 ชั่วโมง 40 นาที"), ("2", "การวัด (เวลา)", "10", "8 ชั่วโมง 20 นาที")]
    for c, h in enumerate(headers):
        set_fixed_width(table.cell(0, c), widths[c]); set_cell_shading(table.cell(0, c), NAVY)
        add_labeled_text(table.cell(0, c), "", h, label_color=WHITE, size=9)
        table.cell(0, c).paragraphs[0].runs[0].font.color.rgb = RGBColor.from_string(WHITE)
        table.cell(0, c).paragraphs[0].runs[0].bold = True
        finish_header_cell(table.cell(0, c))
    for r, row in enumerate(data, start=1):
        for c, value in enumerate(row):
            set_fixed_width(table.cell(r, c), widths[c]); set_cell_margins(table.cell(r, c), top=100, bottom=100)
            add_labeled_text(table.cell(r, c), "", value, size=9)
    set_table_geometry(table, widths)
    set_table_borders(table)

    # Static contents over three pages for reliable printing without field updates.
    chunks = [plans[:26], plans[26:52], plans[52:]]
    for ci, chunk in enumerate(chunks):
        doc.add_page_break()
        add_heading(doc, "สารบัญ" if ci == 0 else "สารบัญ (ต่อ)", size=18)
        for plan in chunk:
            p = doc.add_paragraph()
            set_para(p, after=2.4, line=1.0)
            p.paragraph_format.left_indent = Cm(0.2)
            p.paragraph_format.tab_stops.add_tab_stop(Cm(3.2))
            set_run(p.add_run(f"หน่วย {plan['unit']} · แผนที่ {plan['number']}"), 8.6, True, NAVY)
            set_run(p.add_run("\t" + shorten(plan["title"], 80)), 8.6, color=INK)


def add_unit_divider(doc, unit, lead_break=True):
    center = WD_ALIGN_PARAGRAPH.CENTER
    if lead_break:
        doc.add_page_break()
    for _ in range(8):
        set_para(doc.add_paragraph(), after=0)
    p = doc.add_paragraph()
    set_para(p, after=2, align=center)
    set_run(p.add_run("หน่วยการเรียนรู้"), 16, False, MUTED)
    p = doc.add_paragraph()
    set_para(p, after=4, align=center)
    set_run(p.add_run(str(unit)), 88, True, GOLD)
    p = doc.add_paragraph()
    set_para(p, before=0, after=10, align=center)
    set_run(p.add_run("◆ ◆ ◆"), 13, True, GOLD)
    name = "จำนวนนับ และการบวก การลบ การคูณ การหาร" if unit == 1 else "การวัด (เวลา)"
    p = doc.add_paragraph()
    set_para(p, after=10, align=center)
    set_run(p.add_run(name), 21, True, NAVY)
    count = 62 if unit == 1 else 10
    p = doc.add_paragraph()
    set_para(p, align=center)
    set_run(p.add_run(f"จำนวน {count} แผน   •   แผนละ 50 นาที"), 12, color=MUTED)


def add_appendices(doc, lead_break=True):
    if lead_break:
        doc.add_page_break()
    add_heading(doc, "แบบบันทึกหลังการจัดการเรียนรู้", size=18)
    fields = [
        "แผนการจัดการเรียนรู้ที่ .......... เรื่อง ................................................................................................",
        "วันที่สอน ........................................ ชั้น ป.4/........ จำนวนนักเรียน ............ คน",
        "1. ผลการเรียนรู้ที่เกิดขึ้นกับผู้เรียน\n........................................................................................................................................................\n........................................................................................................................................................",
        "2. ปัญหาและอุปสรรค\n........................................................................................................................................................\n........................................................................................................................................................",
        "3. แนวทางแก้ไข/ข้อเสนอแนะ\n........................................................................................................................................................\n........................................................................................................................................................",
        "4. ผู้เรียนที่ควรได้รับการส่งเสริมหรือช่วยเหลือเพิ่มเติม\n........................................................................................................................................................\n........................................................................................................................................................",
    ]
    for text in fields:
        p = doc.add_paragraph(); set_para(p, after=9, line=1.2); set_run(p.add_run(text), 10)
    p = doc.add_paragraph(); set_para(p, before=12, align=WD_ALIGN_PARAGRAPH.RIGHT); set_run(p.add_run("ลงชื่อ ........................................................ ผู้สอน\n(........................................................)"), 10)

    doc.add_page_break()
    add_heading(doc, "แบบประเมินการปฏิบัติงานทางคณิตศาสตร์", size=18)
    table = doc.add_table(rows=11, cols=5); table.alignment = WD_TABLE_ALIGNMENT.CENTER; table.autofit = False
    headers = ["เลขที่", "ความถูกต้อง", "การให้เหตุผล", "การสื่อสาร", "ผลการประเมิน"]
    widths = [1.6, 4.2, 4.2, 4.2, 3.0]
    for c,h in enumerate(headers):
        set_fixed_width(table.cell(0,c), widths[c]); set_cell_shading(table.cell(0,c), NAVY); add_labeled_text(table.cell(0,c), "", h, label_color=WHITE, size=8.7); table.cell(0,c).paragraphs[0].runs[0].font.color.rgb=RGBColor.from_string(WHITE); table.cell(0,c).paragraphs[0].runs[0].bold=True; finish_header_cell(table.cell(0,c))
    for r in range(1,11):
        for c in range(5):
            set_fixed_width(table.cell(r,c), widths[c]); set_cell_margins(table.cell(r,c), top=150, bottom=150); add_labeled_text(table.cell(r,c), "", str(r) if c==0 else "", size=9)
    set_table_geometry(table, widths)
    set_table_borders(table)
    p=doc.add_paragraph(); set_para(p,before=7,line=1.15); set_run(p.add_run("เกณฑ์: 3 = ดี  |  2 = พอใช้  |  1 = ควรปรับปรุง  •  ผ่านเกณฑ์เมื่อได้ระดับพอใช้ขึ้นไป"),9,color=MUTED)

    doc.add_page_break()
    add_heading(doc, "แบบประเมินคุณลักษณะอันพึงประสงค์", size=18)
    table = doc.add_table(rows=11, cols=5); table.alignment = WD_TABLE_ALIGNMENT.CENTER; table.autofit = False
    headers = ["เลขที่", "ใฝ่เรียนรู้", "ซื่อสัตย์สุจริต", "มุ่งมั่นในการทำงาน", "ผลการประเมิน"]
    for c,h in enumerate(headers):
        set_fixed_width(table.cell(0,c), widths[c]); set_cell_shading(table.cell(0,c), GOLD); add_labeled_text(table.cell(0,c), "", h, label_color=WHITE, size=8.5); table.cell(0,c).paragraphs[0].runs[0].font.color.rgb=RGBColor.from_string(WHITE); table.cell(0,c).paragraphs[0].runs[0].bold=True; finish_header_cell(table.cell(0,c))
    for r in range(1,11):
        for c in range(5):
            set_fixed_width(table.cell(r,c), widths[c]); set_cell_margins(table.cell(r,c), top=150, bottom=150); add_labeled_text(table.cell(r,c), "", str(r) if c==0 else "", size=9)
    set_table_geometry(table, widths)
    set_table_borders(table)
    p=doc.add_paragraph(); set_para(p,before=7,line=1.15); set_run(p.add_run("เกณฑ์: 3 = แสดงพฤติกรรมสม่ำเสมอ  |  2 = แสดงพฤติกรรมบางครั้ง  |  1 = ควรส่งเสริมเพิ่มเติม"),9,color=MUTED)


def set_page_start(doc, start):
    """Force the section's page numbering to begin at `start` (for chunked export)."""
    sec_pr = doc.sections[0]._sectPr
    pg = sec_pr.find(qn("w:pgNumType"))
    if pg is None:
        pg = OxmlElement("w:pgNumType")
        sec_pr.append(pg)
    pg.set(qn("w:start"), str(start))


def configure_doc(doc, first_page_special=True, page_start=None):
    sec = doc.sections[0]
    sec.page_width = Cm(21.0); sec.page_height = Cm(29.7)
    sec.top_margin = Cm(1.35); sec.bottom_margin = Cm(1.35); sec.left_margin = Cm(1.55); sec.right_margin = Cm(1.55)
    sec.header_distance = Cm(0.65); sec.footer_distance = Cm(0.65)
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = FONT; normal.font.size = Pt(9)
    normal._element.rPr.rFonts.set(qn("w:ascii"), FONT); normal._element.rPr.rFonts.set(qn("w:hAnsi"), FONT); normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)

    def fill_header_footer(target):
        p = target.header.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        set_para(p, after=0)
        set_run(p.add_run("แผนการจัดการเรียนรู้ คณิตศาสตร์ ป.4 • ภาคเรียนที่ 1 ปีการศึกษา 2569"), 7.5, color=MUTED)
        fp = target.footer.paragraphs[0]
        set_para(fp, after=0)
        add_page_number(fp)

    fill_header_footer(sec)
    sec.different_first_page_header_footer = first_page_special
    if not first_page_special:
        # ensure the first physical page of a plan chunk also shows the running header/number
        sec.first_page_header.is_linked_to_previous = True
        sec.first_page_footer.is_linked_to_previous = True
    if page_start is not None:
        set_page_start(doc, page_start)


def main():
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    pdf = pymupdf.open(SOURCE)
    plans = detect_plans(pdf)
    doc = Document()
    configure_doc(doc)
    add_cover(doc)
    add_front_matter(doc, plans)
    add_unit_divider(doc, 1)
    for i, plan in enumerate([p for p in plans if p["unit"] == 1]):
        doc.add_page_break()
        add_plan_page(doc, plan)
    add_unit_divider(doc, 2)
    for plan in [p for p in plans if p["unit"] == 2]:
        doc.add_page_break()
        add_plan_page(doc, plan)
    add_appendices(doc)
    doc.core_properties.title = "แผนการจัดการเรียนรู้ คณิตศาสตร์ ป.4 ภาคเรียนที่ 1 ปีการศึกษา 2569"
    doc.core_properties.subject = "แผนการสอนฉบับกระชับพร้อมใช้"
    doc.core_properties.author = "โรงเรียนคำไผ่"
    doc.core_properties.keywords = "แผนการสอน, คณิตศาสตร์, ป.4, DLTV"
    doc.save(OUT)
    print(OUT)
    print(f"plans={len(plans)}")


if __name__ == "__main__":
    main()
